from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture("George Banica Autor.png", width=Inches(1.5))

# name | phone number | email address
speak("What is your name?")
name = input("What is your name? ")
speak("Hello " + name + "How are you today?")
speak("What is your phone number?")
phone_number = input("What is your phone number? ")
email_address = input("What is your email address? ")
document.add_paragraph(name + "  |  " + phone_number + "  |  " + email_address)

# about me
document.add_heading("About me")
document.add_paragraph(input("Tell me about yourself? "))

# work experience
document.add_heading("Work Experience")
p = document.add_paragraph()

company = input("Enter company: ")
from_date = input("From Date: ")
to_date = input("To Date: ")

p.add_run(company + "  ").bold = True
p.add_run(from_date + "-" + to_date + "\n").italic = True

experience_details = input("Describe your experience at " + company + ": ")

p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input("Do you have more experiences? Yes or No: ")
    if has_more_experiences.lower() == "yes":
        p = document.add_paragraph()

        company = input("Enter company: ")
        from_date = input("From Date: ")
        to_date = input("To Date: ")

        p.add_run(company + "  ").bold = True
        p.add_run(from_date + "-" + to_date + "\n").italic = True

        experience_details = input("Describe your experience at " + company + ": ")

        p.add_run(experience_details)
    else:
        break


# skills
document.add_heading("Skills")
skill = input("Input skill: ")
p = document.add_paragraph(skill)
p.style = "List Bullet"

while True:
    has_more_skills = input("Do you have more skills? Yes or No: ")
    if has_more_skills.lower() == "yes":
        skill = input("Input skill: ")
        p = document.add_paragraph(skill)
        p.style = "List Bullet"
    else:
        break


# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "-- CV generated using Amigoscode and Intuit Quickbookscode. --".center(150)


document.save("cv.docx")



"""class Person:
    def __init__(self, name, age):
        self.firstParam = name
        self.secondParam = age

    def walk(self):
        print(self.firstParam + " is walking.")

    def speak(self):
        print("Hello, my name is " + self.firstParam + " and I'm " + str(self.secondParam) + " years old.")


john = Person("John", 22)
john.speak()
john.walk()

miriam = Person("Miriam", 18)
miriam.speak()
miriam.walk()


class Fruit:
    def __init__(self, name, color, weight, size="medium"):
        self.name = name
        self.color = color
        self.weight = weight
        self.size = size

    def falls(self):
        print("The " + self.color + " " + self.name + " is falling from the tree because it's heavy. \nIt weighs " + str(self.weight) + " grams and it's " + self.size + "!\n")

apple = Fruit("apple", "red", 500, "big")
apple.falls()

banana = Fruit("banana", "yellow", 250, "small")
banana.falls()
"""    


"""age1 = 15
age2 = 18
age3 = 8

def check_age(user_age):
    if user_age < 10:
        print("Child")
    elif 10 <= user_age <= 17:
        print("Teenager")
    else:
        print("Adult")

check_age(age1)
check_age(age2)
check_age(age3)
"""


"""number = 0
while number <= 10:
    print(number)
    number += 1
print("The while loop ended because the number is: " + str(number))
"""


"""car_brands = ["bmw", "tesla", "Alfa RomEO", "lexus"]
for car in car_brands:
    if car == "bmw":
        print(car.upper())
    else:
        print(car.capitalize())

# name = "Victor"
"""

"""print(len(car_brands))
print(car_brands)
print(car_brands[2])
print(car_brands[0])
print(car_brands[1])
print(car_brands[3])
print()
"""

"""" first_name = "jamila"
surname = "Smith"
full_name = first_name.capitalize() + " " + surname

print(full_name)
print(len(first_name))
print(len(surname))
print(len(full_name))
print(full_name.endswith("th"))
print(full_name.endswith("ba"))
"""

